"""
test_loaders.py — Phase 1, Chunk 1 tests

WHY THIS EXISTS:
loaders.py is the foundation everything else (chunking, embeddings,
citations) gets built on top of. If page/paragraph numbering is wrong
here, every citation DocMind ever shows will quietly point to the wrong
place — and nothing will crash to tell you. These tests exist to catch
that class of silent bug before it propagates.

The most important test in this file is test_load_docx_preserves_true_position,
which is a regression test for a real bug we found: an earlier version of
load_docx used a counter that only incremented on non-empty paragraphs,
so skipped blank paragraphs silently shifted every later position number.
"""

import pytest
from pathlib import Path
from unittest.mock import patch
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "app" / "ingestion"))
from loaders import load_pdf, load_docx


# ---------------------------------------------------------------------------
# Fixtures — build known-structure test files so we know the "correct"
# answer before running anything, rather than just eyeballing output.
# ---------------------------------------------------------------------------

@pytest.fixture
def simple_docx(tmp_path) -> str:
    """A DOCX with no blanks — the easy case."""
    path = tmp_path / "simple.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(path))
    return str(path)


@pytest.fixture
def docx_with_gaps(tmp_path) -> str:
    """
    A DOCX with blank paragraphs deliberately scattered through it.
    True positions of real content: 1, 3, 6.
    Blanks at: 2, 4, 5.
    This is the file that exposes the indexing bug if it's reintroduced.
    """
    path = tmp_path / "gaps.docx"
    doc = DocxDocument()
    doc.add_paragraph("Paragraph A - has content")   # true position 1
    doc.add_paragraph("")                              # true position 2 - blank
    doc.add_paragraph("Paragraph C - has content")    # true position 3
    doc.add_paragraph("")                              # true position 4 - blank
    doc.add_paragraph("")                              # true position 5 - blank
    doc.add_paragraph("Paragraph F - has content")    # true position 6
    doc.save(str(path))
    return str(path)


@pytest.fixture
def simple_pdf(tmp_path) -> str:
    """A 2-page PDF with simple text content on each page."""
    path = tmp_path / "simple.pdf"
    c = canvas.Canvas(str(path))
    c.drawString(100, 750, "Page one content")
    c.showPage()
    c.drawString(100, 750, "Page two content")
    c.showPage()
    c.save()
    return str(path)


@pytest.fixture
def pdf_with_blank_page(tmp_path) -> str:
    """
    A 3-page PDF where the middle page is genuinely blank (no text drawn).
    True positions of real content: 1, 3. Page 2 should be skipped.
    """
    path = tmp_path / "blank_page.pdf"
    c = canvas.Canvas(str(path))
    c.drawString(100, 750, "Page one content")
    c.showPage()
    # page two intentionally left blank — no drawString call
    c.showPage()
    c.drawString(100, 750, "Page three content")
    c.showPage()
    c.save()
    return str(path)


@pytest.fixture
def empty_docx(tmp_path) -> str:
    """
    A DOCX with only blank paragraphs — no real content anywhere.
    Simulates a realistic edge case: a user uploads a doc that's just
    whitespace, empty headers, or formatting with no actual text.
    Should return [] cleanly, not error or return malformed entries.
    """
    path = tmp_path / "empty.docx"
    doc = DocxDocument()
    doc.add_paragraph("")
    doc.add_paragraph("   ")  # whitespace-only, not just empty string
    doc.add_paragraph("")
    doc.save(str(path))
    return str(path)


@pytest.fixture
def blank_pdf(tmp_path) -> str:
    """
    A PDF with pages but zero extractable text (nothing drawn on any page).
    Simulates a scanned PDF with no OCR layer — a realistic real-world
    case for a document upload feature. Should return [] cleanly.
    """
    path = tmp_path / "blank.pdf"
    c = canvas.Canvas(str(path))
    c.showPage()
    c.showPage()
    c.save()
    return str(path)


@pytest.fixture
def pdf_with_image_page(tmp_path) -> str:
    """
    A 2-page PDF where page 1 has real embedded text and page 2 has none
    (simulates a scanned/image-only page). Used to test OCR fallback:
    pypdf returns empty for page 2, so load_pdf should call _ocr_pdf_page
    for that page and include the OCR result in the output.

    We can't create a genuinely scanned PDF with reportlab — it only
    produces text-based PDFs. So page 2 is simply left blank here, and
    tests that use this fixture mock _ocr_pdf_page to control what "OCR"
    returns, isolating the load_pdf routing logic from Tesseract itself.
    """
    path = tmp_path / "image_page.pdf"
    c = canvas.Canvas(str(path))
    c.drawString(100, 750, "Page one real text")
    c.showPage()
    # Page 2 has no drawn text — pypdf will return empty, triggering OCR
    c.showPage()
    c.save()
    return str(path)


# ---------------------------------------------------------------------------
# load_pdf tests
# ---------------------------------------------------------------------------

def test_load_pdf_returns_correct_page_count(simple_pdf):
    result = load_pdf(simple_pdf)
    assert len(result) == 2


def test_load_pdf_preserves_text_content(simple_pdf):
    result = load_pdf(simple_pdf)
    assert "Page one content" in result[0]["text"]
    assert "Page two content" in result[1]["text"]


def test_load_pdf_sets_locator_type_to_page(simple_pdf):
    result = load_pdf(simple_pdf)
    assert all(entry["locator_type"] == "page" for entry in result)


def test_load_pdf_page_numbers_are_one_indexed(simple_pdf):
    result = load_pdf(simple_pdf)
    assert result[0]["page_number"] == 1
    assert result[1]["page_number"] == 2


def test_load_pdf_skips_blank_pages_but_preserves_true_position(pdf_with_blank_page):
    """
    The middle page has no text. It should be excluded from the output,
    but page 3's page_number should still read 3, not 2 — proving the
    blank page wasn't silently compacted out of the numbering.
    """
    result = load_pdf(pdf_with_blank_page)
    page_numbers = [entry["page_number"] for entry in result]
    assert page_numbers == [1, 3], (
        f"Expected true positions [1, 3] with page 2 skipped, got {page_numbers}"
    )


def test_load_pdf_raises_on_missing_file():
    with pytest.raises(FileNotFoundError):
        load_pdf("/this/path/does/not/exist.pdf")


def test_load_pdf_all_blank_returns_empty_list(blank_pdf):
    """
    A PDF with pages but no extractable text (e.g. a scanned PDF with no
    OCR layer) should return an empty list cleanly — not error, and not
    return malformed entries with empty/whitespace text.
    """
    result = load_pdf(blank_pdf)
    assert result == []


# ---------------------------------------------------------------------------
# load_pdf — OCR fallback tests
#
# WHY MOCKED, NOT REAL TESSERACT: these tests verify the routing logic
# inside load_pdf (does it call OCR when pypdf returns empty? does it use
# the result correctly? does it skip the page when OCR also returns empty?)
# not Tesseract's text recognition quality. Mocking _ocr_pdf_page gives us
# full control over what "OCR" returns so each scenario is exact and fast,
# with no dependency on Tesseract being installed in CI or on other machines.
#
# The mock target is "loaders._ocr_pdf_page" (the name as it exists in the
# loaders module), not a direct import — patching the direct import would
# not intercept the call made from inside load_pdf.
# ---------------------------------------------------------------------------

def test_ocr_fallback_called_when_pypdf_returns_empty(pdf_with_image_page):
    """
    THE CORE OCR ROUTING TEST.

    When pypdf extracts no text from a page, load_pdf must call
    _ocr_pdf_page for that page and include the OCR result in the output.
    This is the entire point of the OCR fallback — without this, any
    scanned page would be silently skipped with no indication to the user
    that content was missed.
    """
    with patch("loaders._ocr_pdf_page", return_value="Scanned page text") as mock_ocr:
        result = load_pdf(pdf_with_image_page)

    # OCR must have been called exactly once — for page 2 only.
    # Page 1 has real text so it must NOT trigger OCR.
    mock_ocr.assert_called_once()

    # Both pages must appear in output: page 1 from pypdf, page 2 from OCR
    assert len(result) == 2
    assert result[1]["page_number"] == 2
    assert result[1]["text"] == "Scanned page text"
    assert result[1]["locator_type"] == "page"


def test_ocr_fallback_called_with_correct_page_number(pdf_with_image_page):
    """
    _ocr_pdf_page must be called with the correct 1-indexed page number.
    If the wrong page number is passed, OCR would extract the wrong page
    from the PDF — a silent error where the text comes back but points
    to the wrong source.
    """
    with patch("loaders._ocr_pdf_page", return_value="Scanned page text") as mock_ocr:
        load_pdf(pdf_with_image_page)

    # Page 2 is the image page — OCR must be called with page_number=2
    args, _ = mock_ocr.call_args
    assert args[1] == 2, (
        f"Expected _ocr_pdf_page to be called with page_number=2, got {args[1]}. "
        f"A wrong page number means OCR extracts the wrong page from the PDF."
    )


def test_ocr_result_skipped_when_ocr_also_returns_empty(pdf_with_image_page):
    """
    If pypdf returns empty AND OCR also returns empty (e.g. a truly blank
    scanned page, or a decorative image with no text), the page must be
    skipped entirely — same behavior as a blank page in a text-based PDF.
    No empty entry should appear in the output.
    """
    with patch("loaders._ocr_pdf_page", return_value=""):
        result = load_pdf(pdf_with_image_page)

    # Only page 1 (real text) should appear; page 2 (empty OCR) skipped
    assert len(result) == 1
    assert result[0]["page_number"] == 1


def test_ocr_not_called_for_pages_with_real_text(simple_pdf):
    """
    Pages that already have extractable text must NOT trigger OCR —
    OCR is the fallback, not the default. Calling OCR on text-based pages
    would be wasteful (pdf2image + Tesseract is slow) and could produce
    slightly different text than pypdf's direct extraction.
    """
    with patch("loaders._ocr_pdf_page") as mock_ocr:
        result = load_pdf(simple_pdf)

    # simple_pdf has real text on both pages — OCR must never be called
    mock_ocr.assert_not_called()
    assert len(result) == 2


def test_ocr_page_number_preserved_in_mixed_pdf(pdf_with_image_page):
    """
    In a PDF that mixes text-based and image-based pages, the true page
    number of the OCR'd page must be preserved correctly — not compacted
    as if the image page didn't exist.

    This is the same class of traceability concern as the DOCX blank-
    paragraph indexing bug: if page numbering gets shifted by the OCR
    path, citations point to the wrong page with no visible symptom.
    """
    with patch("loaders._ocr_pdf_page", return_value="OCR recovered text"):
        result = load_pdf(pdf_with_image_page)

    page_numbers = [r["page_number"] for r in result]
    assert page_numbers == [1, 2], (
        f"Expected page numbers [1, 2] for a 2-page mixed PDF, got {page_numbers}. "
        f"The OCR path must preserve true page position, not re-number pages."
    )


def test_ocr_failure_does_not_crash_ingestion(pdf_with_image_page):
    """
    If _ocr_pdf_page raises an exception (Tesseract not found, pdf2image
    error, etc.), load_pdf must not crash — it should skip that page and
    continue processing the rest of the document.

    This matters because OCR failures are environment-dependent (Tesseract
    not installed, wrong path, corrupted page image) and should degrade
    gracefully rather than blocking the entire document from being ingested.

    NOTE: _ocr_pdf_page already catches exceptions internally and returns
    "". This test verifies that contract holds — if it ever changes to
    raise instead, load_pdf would break here.
    """
    with patch("loaders._ocr_pdf_page", side_effect=Exception("Tesseract not found")):
        # Should not raise — OCR failure must be handled gracefully
        try:
            result = load_pdf(pdf_with_image_page)
        except Exception as e:
            pytest.fail(
                f"load_pdf raised {type(e).__name__} when _ocr_pdf_page failed: {e}. "
                f"OCR failures must be handled gracefully — the page should be "
                f"skipped, not the whole document ingestion aborted."
            )

    # Only page 1 (real text via pypdf) should appear; page 2 (OCR failed) skipped
    assert len(result) == 1
    assert result[0]["page_number"] == 1


# ---------------------------------------------------------------------------
# load_docx tests
# ---------------------------------------------------------------------------

def test_load_docx_returns_correct_paragraph_count(simple_docx):
    result = load_docx(simple_docx)
    assert len(result) == 2


def test_load_docx_preserves_text_content(simple_docx):
    result = load_docx(simple_docx)
    assert result[0]["text"] == "First paragraph."
    assert result[1]["text"] == "Second paragraph."


def test_load_docx_sets_locator_type_to_paragraph_index(simple_docx):
    result = load_docx(simple_docx)
    assert all(entry["locator_type"] == "paragraph_index" for entry in result)


def test_load_docx_raises_on_missing_file():
    with pytest.raises(FileNotFoundError):
        load_docx("/this/path/does/not/exist.docx")


def test_load_docx_preserves_true_position(docx_with_gaps):
    """
    REGRESSION TEST for the indexing bug found during development.

    An earlier version of load_docx used a counter that only incremented
    on non-empty paragraphs. That made "page_number" a count of "how many
    non-blank paragraphs we've seen so far" rather than the paragraph's
    actual position in the document. The bug was silent — no crash, no
    error — it just meant every citation pointing past the first blank
    paragraph would be wrong by however many blanks preceded it.

    This test uses a file with known gaps (blanks at positions 2, 4, 5)
    and asserts the real content lands at its TRUE positions: 1, 3, 6 —
    not the compacted [1, 2, 3] the buggy version would have produced.
    """
    result = load_docx(docx_with_gaps)
    page_numbers = [entry["page_number"] for entry in result]
    assert page_numbers == [1, 3, 6], (
        f"Expected true positions [1, 3, 6] with blanks at 2/4/5 skipped, "
        f"got {page_numbers}. If this is [1, 2, 3], the dense-counter bug "
        f"has been reintroduced."
    )


def test_load_docx_empty_paragraphs_excluded_from_output(docx_with_gaps):
    """Blank paragraphs should not appear as empty entries in the output."""
    result = load_docx(docx_with_gaps)
    assert all(entry["text"].strip() != "" for entry in result)


def test_load_docx_all_blank_returns_empty_list(empty_docx):
    """
    A DOCX with only blank/whitespace paragraphs (no real content at all)
    should return an empty list cleanly — a realistic case when a user
    uploads a doc that's just formatting or empty sections.
    """
    result = load_docx(empty_docx)
    assert result == []


# ---------------------------------------------------------------------------
# Cross-loader consistency — both loaders must return the same shape so
# downstream code (chunker, citation UI) can treat them uniformly.
# ---------------------------------------------------------------------------

def test_pdf_and_docx_return_same_dict_keys(simple_pdf, simple_docx):
    pdf_result = load_pdf(simple_pdf)
    docx_result = load_docx(simple_docx)
    assert set(pdf_result[0].keys()) == set(docx_result[0].keys()) == {
        "page_number", "locator_type", "text"
    }


def test_pdf_and_docx_return_same_value_types(simple_pdf, simple_docx):
    """
    Same keys isn't enough — downstream code (chunker, citation UI) will
    also break if one loader returns page_number as a string and the
    other as an int, or if text comes back empty/None in either case.
    """
    for result in (load_pdf(simple_pdf), load_docx(simple_docx)):
        for entry in result:
            assert isinstance(entry["page_number"], int)
            assert isinstance(entry["locator_type"], str)
            assert isinstance(entry["text"], str)
            assert entry["text"].strip() != ""